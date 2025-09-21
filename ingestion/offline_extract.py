import os
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import numpy as np
from PIL import Image
import cv2
import pytesseract
import whisper
import torch
from transformers import CLIPProcessor, CLIPModel
import docx
import fitz  # PyMuPDF
import tempfile
from pydub import AudioSegment

@dataclass
class OfflineChunk:
    """Container for offline extracted content"""
    content: str
    chunk_type: str  # 'text', 'image', 'audio', 'table'
    source_file: str
    page_number: Optional[int] = None
    timestamp: Optional[float] = None  # For audio
    section_header: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    image_path: Optional[str] = None
    audio_path: Optional[str] = None
    ocr_text: Optional[str] = None
    transcript: Optional[str] = None

class OfflineExtractor:
    """Extract content from multiple file formats completely offline"""
    
    def __init__(self, raw_dir: str = "data/raw", extracted_dir: str = "data/extracted"):
        self.raw_dir = Path(raw_dir)
        self.extracted_dir = Path(extracted_dir)
        self.extracted_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize models
        self._init_models()
        
        # Supported formats
        self.supported_formats = {
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'],
            'audio': ['.wav', '.mp3', '.m4a', '.flac', '.ogg', '.aac']
        }
    
    def _init_models(self):
        """Initialize AI models for processing"""
        try:
            print("Loading CLIP model...")
            self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
            self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
            
            print("Loading Whisper model...")
            self.whisper_model = whisper.load_model("base")
            
            print("All models loaded successfully")
            
        except Exception as e:
            print(f"Failed to initialize models: {e}")
            self.clip_model = None
            self.clip_processor = None
            self.whisper_model = None
    
    def get_file_hash(self, file_path: Path) -> str:
        """Generate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()[:16]
    
    def detect_file_type(self, file_path: Path) -> str:
        """Detect file type based on extension"""
        ext = file_path.suffix.lower()
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        return 'unknown'
    
    def extract_pdf_content(self, pdf_path: Path) -> List[OfflineChunk]:
        """Extract PDF content using PyMuPDF"""
        try:
            chunks = []
            file_hash = self.get_file_hash(pdf_path)
            
            # Create output directory for this PDF
            output_dir = self.extracted_dir / f"{pdf_path.stem}_{file_hash}"
            output_dir.mkdir(exist_ok=True)
            image_path_dir = output_dir / "images"
            image_path_dir.mkdir(exist_ok=True)
            
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    chunk = OfflineChunk(
                        content=text,
                        chunk_type="text",
                        source_file=pdf_path.name,
                        page_number=page_num + 1,
                        metadata={
                            "document_id": file_hash,
                            "source_file": pdf_path.name,
                        }
                    )
                    chunks.append(chunk)
                
                # Extract images
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_filename = f"page_{page_num + 1}_image_{img_index}.png"
                            img_path = image_path_dir / img_filename
                            
                            with open(img_path, "wb") as f:
                                f.write(img_data)
                            
                            # Perform OCR
                            ocr_text = ""
                            try:
                                image = Image.open(img_path)
                                ocr_text = pytesseract.image_to_string(image).strip()
                            except Exception as e:
                                print(f"OCR failed for {img_path}: {e}")
                            
                            chunk = OfflineChunk(
                                content=f"Image from page {page_num + 1}" + (f" - OCR: {ocr_text}" if ocr_text else ""),
                                chunk_type="image",
                                source_file=pdf_path.name,
                                page_number=page_num + 1,
                                image_path=str(img_path),
                                ocr_text=ocr_text,
                                metadata={
                                    "document_id": file_hash,
                                    "source_file": pdf_path.name,
                                }
                            )
                            chunks.append(chunk)
                        
                        pix = None
                    except Exception as e:
                        print(f"Error extracting image: {e}")
                        continue
            
            doc.close()
            print(f"Extracted {len(chunks)} chunks from PDF: {pdf_path}")
            return chunks
            
        except Exception as e:
            print(f"Failed to extract PDF content: {e}")
            return []
    
    def extract_docx_content(self, docx_path: Path) -> List[OfflineChunk]:
        """Extract content from DOCX files"""
        try:
            doc = docx.Document(docx_path)
            chunks = []
            
            file_hash = self.get_file_hash(docx_path)
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    chunk = OfflineChunk(
                        content=paragraph.text,
                        chunk_type='text',
                        source_file=docx_path.name,
                        metadata={
                            'document_id': file_hash,
                            'source_file': docx_path.name,
                            'paragraph_index': i
                        }
                    )
                    chunks.append(chunk)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                
                if table_text:
                    chunk = OfflineChunk(
                        content='\n'.join(table_text),
                        chunk_type='table',
                        source_file=docx_path.name,
                        metadata={
                            'document_id': file_hash,
                            'source_file': docx_path.name,
                            'table_index': i
                        }
                    )
                    chunks.append(chunk)
            
            print(f"Extracted {len(chunks)} chunks from DOCX: {docx_path}")
            return chunks
            
        except Exception as e:
            print(f"Failed to extract DOCX content: {e}")
            return []
    
    def extract_image_content(self, image_path: Path) -> List[OfflineChunk]:
        """Extract content from image files using OCR"""
        try:
            chunks = []
            file_hash = self.get_file_hash(image_path)
            
            # Load image
            image = Image.open(image_path)
            
            # Perform OCR
            ocr_text = ""
            try:
                ocr_text = pytesseract.image_to_string(image).strip()
            except Exception as e:
                print(f"OCR failed for {image_path}: {e}")
            
            # Create image chunk
            chunk = OfflineChunk(
                content=f"Image: {image_path.name}" + (f" - OCR Text: {ocr_text}" if ocr_text else ""),
                chunk_type='image',
                source_file=image_path.name,
                image_path=str(image_path),
                ocr_text=ocr_text,
                metadata={
                    'document_id': file_hash,
                    'source_file': image_path.name,
                    'image_size': image.size,
                    'image_mode': image.mode
                }
            )
            chunks.append(chunk)
            
            print(f"Extracted content from image: {image_path}")
            return chunks
            
        except Exception as e:
            print(f"Failed to extract image content: {e}")
            return []
    
    def extract_audio_content(self, audio_path: Path) -> List[OfflineChunk]:
        """Extract content from audio files using Whisper"""
        try:
            chunks = []
            file_hash = self.get_file_hash(audio_path)
            
            # Convert audio to WAV if needed
            temp_wav = None
            if audio_path.suffix.lower() != '.wav':
                temp_wav = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
                try:
                    audio = AudioSegment.from_file(audio_path)
                    audio.export(temp_wav.name, format='wav')
                    wav_path = temp_wav.name
                except Exception as e:
                    print(f"Audio conversion failed: {e}")
                    return []
            else:
                wav_path = str(audio_path)
            
            try:
                # Use Whisper for transcription
                if self.whisper_model:
                    result = self.whisper_model.transcribe(wav_path)
                    transcript = result['text'].strip()
                    
                    # Create segments if available
                    if 'segments' in result and result['segments']:
                        for segment in result['segments']:
                            chunk = OfflineChunk(
                                content=segment['text'].strip(),
                                chunk_type='audio',
                                source_file=audio_path.name,
                                timestamp=segment['start'],
                                audio_path=str(audio_path),
                                transcript=segment['text'].strip(),
                                metadata={
                                    'document_id': file_hash,
                                    'source_file': audio_path.name,
                                    'start_time': segment['start'],
                                    'end_time': segment['end'],
                                    'duration': segment['end'] - segment['start']
                                }
                            )
                            chunks.append(chunk)
                    else:
                        # Single chunk for entire audio
                        chunk = OfflineChunk(
                            content=transcript,
                            chunk_type='audio',
                            source_file=audio_path.name,
                            audio_path=str(audio_path),
                            transcript=transcript,
                            metadata={
                                'document_id': file_hash,
                                'source_file': audio_path.name,
                                'duration': result.get('duration', 0)
                            }
                        )
                        chunks.append(chunk)
                
            finally:
                # Clean up temporary file
                if temp_wav:
                    try:
                        os.unlink(temp_wav.name)
                    except:
                        pass
            
            print(f"Extracted {len(chunks)} audio segments from: {audio_path}")
            return chunks
            
        except Exception as e:
            print(f"Failed to extract audio content: {e}")
            return []
    
    def extract_file_content(self, file_path: Path) -> List[OfflineChunk]:
        """Extract content from any supported file type"""
        file_type = self.detect_file_type(file_path)
        
        if file_type == 'pdf':
            return self.extract_pdf_content(file_path)
        elif file_type == 'docx':
            return self.extract_docx_content(file_path)
        elif file_type == 'image':
            return self.extract_image_content(file_path)
        elif file_type == 'audio':
            return self.extract_audio_content(file_path)
        else:
            print(f"Unsupported file type: {file_path}")
            return []
    
    def extract_all_files(self) -> Dict[str, List[OfflineChunk]]:
        """Extract content from all supported files in raw directory"""
        all_files = []
        
        # Find all supported files
        for file_type, extensions in self.supported_formats.items():
            for ext in extensions:
                all_files.extend(self.raw_dir.glob(f"*{ext}"))
        
        if not all_files:
            print(f"No supported files found in {self.raw_dir}")
            return {}
        
        print(f"Found {len(all_files)} files to process")
        all_chunks = {}
        
        for file_path in all_files:
            try:
                chunks = self.extract_file_content(file_path)
                if chunks:
                    all_chunks[str(file_path)] = chunks
                    print(f"Successfully processed {file_path}: {len(chunks)} chunks")
                else:
                    print(f"No content extracted from {file_path}")
            except Exception as e:
                print(f"Failed to process {file_path}: {e}")
        
        return all_chunks