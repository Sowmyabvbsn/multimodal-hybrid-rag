import os
import json
import hashlib
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
import numpy as np
from PIL import Image
import cv2
import pytesseract
import whisper
import speech_recognition as sr
from pydub import AudioSegment
import librosa
import torch
from transformers import CLIPProcessor, CLIPModel
import docx
from loguru import logger
import base64
from io import BytesIO

# Import existing PDF extractor
from ingestion.extract import PDFExtractor, ExtractedChunk

@dataclass
class MultimodalChunk:
    """Container for multimodal extracted content"""
    content: str
    chunk_type: str  # 'text', 'image', 'audio', 'table'
    source_file: str
    page_number: Optional[int] = None
    timestamp: Optional[float] = None  # For audio
    section_header: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    image_path: Optional[str] = None
    audio_path: Optional[str] = None
    ocr_text: Optional[str] = None
    transcript: Optional[str] = None
    embeddings: Optional[Dict[str, np.ndarray]] = None

class MultimodalExtractor:
    """Extract content from multiple file formats: PDF, DOCX, images, audio"""
    
    def __init__(self, raw_dir: str = "data/raw", extracted_dir: str = "data/extracted"):
        self.raw_dir = Path(raw_dir)
        self.extracted_dir = Path(extracted_dir)
        self.extracted_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize PDF extractor
        self.pdf_extractor = PDFExtractor(raw_dir, extracted_dir)
        
        # Initialize models
        self._init_models()
        
        # Supported formats
        self.supported_formats = {
            'pdf': ['.pdf'],
            'docx': ['.docx', '.doc'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp'],
            'audio': ['.wav', '.mp3', '.m4a', '.flac', '.ogg', '.aac']
        }
    
    def _init_models(self):
        """Initialize AI models for processing"""
        try:
            # Initialize CLIP for image embeddings
            logger.info("Loading CLIP model...")
            self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
            self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
            
            # Initialize Whisper for audio transcription
            logger.info("Loading Whisper model...")
            self.whisper_model = whisper.load_model("base")
            
            # Initialize speech recognition as backup
            self.sr_recognizer = sr.Recognizer()
            
            logger.info("All models loaded successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize models: {e}")
            self.clip_model = None
            self.clip_processor = None
            self.whisper_model = None
            self.sr_recognizer = None
    
    def get_file_hash(self, file_path: Path) -> str:
        """Generate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()[:16]
    
    def detect_file_type(self, file_path: Path) -> str:
        """Detect file type based on extension"""
        ext = file_path.suffix.lower()
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        return 'unknown'
    
    def extract_pdf_content(self, pdf_path: Path) -> List[MultimodalChunk]:
        """Extract PDF content using existing extractor"""
        try:
            chunks_data = self.pdf_extractor.extract_pdf_content(pdf_path)
            multimodal_chunks = []
            
            for page_data in chunks_data:
                # Text content
                if page_data.get('text', '').strip():
                    chunk = MultimodalChunk(
                        content=page_data['text'],
                        chunk_type='text',
                        source_file=pdf_path.name,
                        page_number=page_data['metadata']['page_number'],
                        metadata=page_data['metadata']
                    )
                    multimodal_chunks.append(chunk)
                
                # Image content
                for image_data in page_data.get('image', []):
                    chunk = MultimodalChunk(
                        content=image_data.get('content', ''),
                        chunk_type='image',
                        source_file=pdf_path.name,
                        page_number=page_data['metadata']['page_number'],
                        image_path=image_data.get('image_path'),
                        section_header=image_data.get('section_header'),
                        metadata=page_data['metadata']
                    )
                    multimodal_chunks.append(chunk)
                
                # Table content
                for table_data in page_data.get('table', []):
                    chunk = MultimodalChunk(
                        content=table_data.get('content', ''),
                        chunk_type='table',
                        source_file=pdf_path.name,
                        page_number=page_data['metadata']['page_number'],
                        section_header=table_data.get('section_header'),
                        metadata=page_data['metadata']
                    )
                    multimodal_chunks.append(chunk)
            
            return multimodal_chunks
            
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            return []
    
    def extract_docx_content(self, docx_path: Path) -> List[MultimodalChunk]:
        """Extract content from DOCX files"""
        try:
            doc = docx.Document(docx_path)
            chunks = []
            
            file_hash = self.get_file_hash(docx_path)
            
            # Extract paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    chunk = MultimodalChunk(
                        content=paragraph.text,
                        chunk_type='text',
                        source_file=docx_path.name,
                        metadata={
                            'document_id': file_hash,
                            'source_file': docx_path.name,
                            'paragraph_index': i
                        }
                    )
                    chunks.append(chunk)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                
                if table_text:
                    chunk = MultimodalChunk(
                        content='\n'.join(table_text),
                        chunk_type='table',
                        source_file=docx_path.name,
                        metadata={
                            'document_id': file_hash,
                            'source_file': docx_path.name,
                            'table_index': i
                        }
                    )
                    chunks.append(chunk)
            
            logger.info(f"Extracted {len(chunks)} chunks from DOCX: {docx_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {e}")
            return []
    
    def extract_image_content(self, image_path: Path) -> List[MultimodalChunk]:
        """Extract content from image files using OCR and CLIP"""
        try:
            chunks = []
            file_hash = self.get_file_hash(image_path)
            
            # Load image
            image = Image.open(image_path)
            
            # Perform OCR
            ocr_text = ""
            try:
                ocr_text = pytesseract.image_to_string(image).strip()
            except Exception as e:
                logger.warning(f"OCR failed for {image_path}: {e}")
            
            # Create image chunk
            chunk = MultimodalChunk(
                content=f"Image: {image_path.name}" + (f" - OCR Text: {ocr_text}" if ocr_text else ""),
                chunk_type='image',
                source_file=image_path.name,
                image_path=str(image_path),
                ocr_text=ocr_text,
                metadata={
                    'document_id': file_hash,
                    'source_file': image_path.name,
                    'image_size': image.size,
                    'image_mode': image.mode
                }
            )
            chunks.append(chunk)
            
            logger.info(f"Extracted content from image: {image_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to extract image content: {e}")
            return []
    
    def extract_audio_content(self, audio_path: Path) -> List[MultimodalChunk]:
        """Extract content from audio files using speech recognition"""
        try:
            chunks = []
            file_hash = self.get_file_hash(audio_path)
            
            # Convert audio to WAV if needed
            temp_wav = None
            if audio_path.suffix.lower() != '.wav':
                temp_wav = tempfile.NamedTemporaryFile(suffix='.wav', delete=False)
                audio = AudioSegment.from_file(audio_path)
                audio.export(temp_wav.name, format='wav')
                wav_path = temp_wav.name
            else:
                wav_path = str(audio_path)
            
            try:
                # Use Whisper for transcription
                if self.whisper_model:
                    result = self.whisper_model.transcribe(wav_path)
                    transcript = result['text'].strip()
                    
                    # Create segments if available
                    if 'segments' in result:
                        for segment in result['segments']:
                            chunk = MultimodalChunk(
                                content=segment['text'].strip(),
                                chunk_type='audio',
                                source_file=audio_path.name,
                                timestamp=segment['start'],
                                audio_path=str(audio_path),
                                transcript=segment['text'].strip(),
                                metadata={
                                    'document_id': file_hash,
                                    'source_file': audio_path.name,
                                    'start_time': segment['start'],
                                    'end_time': segment['end'],
                                    'duration': segment['end'] - segment['start']
                                }
                            )
                            chunks.append(chunk)
                    else:
                        # Single chunk for entire audio
                        chunk = MultimodalChunk(
                            content=transcript,
                            chunk_type='audio',
                            source_file=audio_path.name,
                            audio_path=str(audio_path),
                            transcript=transcript,
                            metadata={
                                'document_id': file_hash,
                                'source_file': audio_path.name,
                                'duration': librosa.get_duration(filename=wav_path)
                            }
                        )
                        chunks.append(chunk)
                
                else:
                    # Fallback to speech_recognition
                    with sr.AudioFile(wav_path) as source:
                        audio_data = self.sr_recognizer.record(source)
                        transcript = self.sr_recognizer.recognize_google(audio_data)
                        
                        chunk = MultimodalChunk(
                            content=transcript,
                            chunk_type='audio',
                            source_file=audio_path.name,
                            audio_path=str(audio_path),
                            transcript=transcript,
                            metadata={
                                'document_id': file_hash,
                                'source_file': audio_path.name,
                                'duration': librosa.get_duration(filename=wav_path)
                            }
                        )
                        chunks.append(chunk)
                
            finally:
                # Clean up temporary file
                if temp_wav:
                    os.unlink(temp_wav.name)
            
            logger.info(f"Extracted {len(chunks)} audio segments from: {audio_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to extract audio content: {e}")
            return []
    
    def extract_file_content(self, file_path: Path) -> List[MultimodalChunk]:
        """Extract content from any supported file type"""
        file_type = self.detect_file_type(file_path)
        
        if file_type == 'pdf':
            return self.extract_pdf_content(file_path)
        elif file_type == 'docx':
            return self.extract_docx_content(file_path)
        elif file_type == 'image':
            return self.extract_image_content(file_path)
        elif file_type == 'audio':
            return self.extract_audio_content(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_path}")
            return []
    
    def extract_all_files(self) -> Dict[str, List[MultimodalChunk]]:
        """Extract content from all supported files in raw directory"""
        all_files = []
        
        # Find all supported files
        for file_type, extensions in self.supported_formats.items():
            for ext in extensions:
                all_files.extend(self.raw_dir.glob(f"*{ext}"))
        
        if not all_files:
            logger.warning(f"No supported files found in {self.raw_dir}")
            return {}
        
        logger.info(f"Found {len(all_files)} files to process")
        all_chunks = {}
        
        for file_path in all_files:
            try:
                chunks = self.extract_file_content(file_path)
                if chunks:
                    all_chunks[str(file_path)] = chunks
                    logger.info(f"Successfully processed {file_path}: {len(chunks)} chunks")
                else:
                    logger.warning(f"No content extracted from {file_path}")
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
        
        return all_chunks
    
    def save_chunks_metadata(self, all_chunks: Dict[str, List[MultimodalChunk]], output_file: str = "multimodal_chunks.json"):
        """Save extracted chunks metadata to JSON file"""
        output_path = self.extracted_dir / output_file
        
        serializable_chunks = {}
        for file_path, chunks in all_chunks.items():
            serializable_chunks[file_path] = []
            for chunk in chunks:
                chunk_data = {
                    'content': chunk.content[:500] + "..." if len(chunk.content) > 500 else chunk.content,
                    'chunk_type': chunk.chunk_type,
                    'source_file': chunk.source_file,
                    'page_number': chunk.page_number,
                    'timestamp': chunk.timestamp,
                    'section_header': chunk.section_header,
                    'metadata': chunk.metadata,
                    'image_path': chunk.image_path,
                    'audio_path': chunk.audio_path,
                    'ocr_text': chunk.ocr_text[:200] + "..." if chunk.ocr_text and len(chunk.ocr_text) > 200 else chunk.ocr_text,
                    'transcript': chunk.transcript[:200] + "..." if chunk.transcript and len(chunk.transcript) > 200 else chunk.transcript
                }
                serializable_chunks[file_path].append(chunk_data)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(serializable_chunks, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved chunks metadata to {output_path}")

def main():
    """Example usage"""
    extractor = MultimodalExtractor()
    all_chunks = extractor.extract_all_files()
    
    # Print summary
    total_chunks = sum(len(chunks) for chunks in all_chunks.values())
    print(f"\nExtracted {total_chunks} chunks from {len(all_chunks)} files:")
    
    for file_path, chunks in all_chunks.items():
        chunk_types = {}
        for chunk in chunks:
            chunk_types[chunk.chunk_type] = chunk_types.get(chunk.chunk_type, 0) + 1
        
        print(f"\n{Path(file_path).name}: {len(chunks)} chunks")
        for chunk_type, count in chunk_types.items():
            print(f"  - {chunk_type}: {count}")
    
    # Save metadata
    extractor.save_chunks_metadata(all_chunks)

if __name__ == "__main__":
    main()